from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks, Request
from datetime import datetime

from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
import os
import uuid
from pdf2docx import Converter
from fastapi import Depends
from fastapi.security import OAuth2PasswordRequestForm
from sqlalchemy.orm import Session
from datetime import timedelta
# `Optional` 73- va 92-qatorlarda ishlatiladi, lekin import qilinmagan edi —
# shu sababdan modul umuman yuklanmasdi va ilova ishga tushmasdi.
from typing import List, Optional
from pypdf import PdfWriter, PdfReader, PageObject
from pdf2image import convert_from_path
import zipfile
import shutil
import subprocess
import io
import json
import base64
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import Color
from reportlab.lib.utils import ImageReader
from PIL import Image as PILImage
import openpyxl
from pydantic import BaseModel
from fastapi import Form
import httpx
from fastapi.responses import RedirectResponse, HTMLResponse
import database
import auth

app = FastAPI(title="UzbekPDF API")

# Allow CORS for local development (and later for production)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"]
)

app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/favicon.ico", include_in_schema=False)
async def favicon():
    return FileResponse("favicon.ico")


TEMP_DIR = "temp"
os.makedirs(TEMP_DIR, exist_ok=True)

from jose import jwt
async def get_optional_user(request: Request, db: Session = Depends(database.get_db)):
    auth_header = request.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        return None
    token = auth_header.split(" ")[1]
    try:
        payload = jwt.decode(token, auth.SECRET_KEY, algorithms=[auth.ALGORITHM])
        username: str = payload.get("sub")
        if username:
            return db.query(database.User).filter(database.User.username == username).first()
    except:
        pass
    return None

async def check_usage_limit(file_size: int, request: Request, db: Session, user: Optional[database.User]):
    if user and user.is_premium:
        return
    
    if file_size > 20971520: # 20MB
        raise HTTPException(status_code=413, detail="Fayl hajmi 20MB dan ko'p. Premiumga o'ting!")
    
    ip = request.client.host
    today = datetime.utcnow().strftime("%Y-%m-%d")
    
    query = db.query(database.Usage).filter(database.Usage.date == today)
    if user:
        usage = query.filter(database.Usage.user_id == user.id).first()
    else:
        usage = query.filter(database.Usage.ip_address == ip, database.Usage.user_id == None).first()
        
    if usage and usage.count >= 3:
        raise HTTPException(status_code=403, detail="Kunlik limit (3 ta fayl) tugadi. Premiumga o'ting!")

async def increment_usage(request: Request, db: Session, user: Optional[database.User]):
    if user and user.is_premium:
        return
        
    ip = request.client.host
    today = datetime.utcnow().strftime("%Y-%m-%d")
    
    query = db.query(database.Usage).filter(database.Usage.date == today)
    if user:
        usage = query.filter(database.Usage.user_id == user.id).first()
    else:
        usage = query.filter(database.Usage.ip_address == ip, database.Usage.user_id == None).first()
        
    if usage:
        usage.count += 1
    else:
        usage = database.Usage(
            user_id=user.id if user else None,
            ip_address=ip,
            count=1,
            date=today
        )
        db.add(usage)
    db.commit()

async def save_upload_file(upload_file: UploadFile, destination: str):
    with open(destination, "wb") as buffer:
        while True:
            chunk = await upload_file.read(1024 * 1024)
            if not chunk: break
            buffer.write(chunk)



def cleanup_files(*file_paths):
    for path in file_paths:
        try:
            if os.path.exists(path):
                os.remove(path)
        except Exception as e:
            print(f"Error deleting file {path}: {e}")

class UserCreate(BaseModel):
    email: str
    full_name: str
    password: str



@app.post("/api/register")
def register_user(user: UserCreate, db: Session = Depends(database.get_db)):
    db_user = db.query(database.User).filter(database.User.email == user.email).first()
    if db_user:
        raise HTTPException(status_code=400, detail="Ushbu email band, boshqa email tanlang")
    
    hashed_password = auth.get_password_hash(user.password)
    new_user = database.User(
        username=user.full_name, # Use full name as username display
        email=user.email, 
        password_hash=hashed_password
    )
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    return {"message": "Muvaffaqiyatli ro'yxatdan o'tdingiz"}



@app.post("/api/login")
def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(database.get_db)):
    # Treat 'username' field in form_data as 'email'
    user = db.query(database.User).filter(database.User.email == form_data.username).first()
    if not user or not auth.verify_password(form_data.password, user.password_hash):
        raise HTTPException(
            status_code=401,
            detail="Noto'g'ri email yoki parol",
            headers={"WWW-Authenticate": "Bearer"},
        )

    access_token_expires = timedelta(minutes=auth.ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = auth.create_access_token(
        data={"sub": user.username}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer", "username": user.username}

# Kalitlar `.env` dan olinadi va repoga tushmaydi. Ilgari ular shu yerda
# ochiq yozilgan edi — ya'ni repoga kirgan har kim Google hisobingiz nomidan
# so'rov yubora olardi. Eski kalit bekor qilinishi kerak, yangisini
# almashtirish esa endi kodga tegmasdan bo'ladi.
GOOGLE_CLIENT_ID = os.environ.get("GOOGLE_CLIENT_ID", "")
GOOGLE_CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET", "")

# Manzil ham sozlamada: lokalda sinaganda `http://localhost:8000` bo'ladi.
# Qotirib qo'yilsa Google kirgan odamni prod saytiga qaytarib yuborardi va
# lokalda kirishni umuman sinab bo'lmasdi.
BASE_URL = os.environ.get("BASE_URL", "https://uzbekpdf.uz").rstrip("/")
GOOGLE_REDIRECT_URI = f"{BASE_URL}/api/auth/google/callback"


@app.get("/api/user/me")
async def get_me(request: Request, db: Session = Depends(database.get_db), user: database.User = Depends(auth.get_current_user)):
    # Get daily usage count
    today = datetime.utcnow().strftime("%Y-%m-%d")
    usage = db.query(database.Usage).filter(
        database.Usage.user_id == user.id,
        database.Usage.date == today
    ).first()
    
    return {
        "username": user.username,
        "is_premium": user.is_premium,
        "avatar_url": user.avatar_url,
        "usage_count": usage.count if usage else 0
    }

@app.get("/api/auth/google/login")
def google_login():
    """Google bilan kirishni boshlaydi.

    Bu funksiyada `@app.get` belgisi umuman yo'q edi — ya'ni u oddiy
    funksiya bo'lib qolgan va manzil 404 qaytargan. Kirish tugmasi
    bosilganda hech narsa bo'lmasligining sababi shu edi.
    """
    if not GOOGLE_CLIENT_ID or not GOOGLE_CLIENT_SECRET:
        # Kalit sozlanmagan bo'lsa Google'ga bo'sh `client_id` bilan
        # yuborishning ma'nosi yo'q: u tushunarsiz xato sahifasini
        # ko'rsatadi. Sababni o'zimiz aytamiz.
        raise HTTPException(
            status_code=503,
            detail="Google bilan kirish hozircha sozlanmagan",
        )

    google_auth_url = (
        "https://accounts.google.com/o/oauth2/v2/auth"
        f"?client_id={GOOGLE_CLIENT_ID}"
        f"&redirect_uri={GOOGLE_REDIRECT_URI}"
        "&response_type=code&scope=openid email profile"
    )
    return RedirectResponse(google_auth_url)

async def get_admin_user(user: database.User = Depends(auth.get_current_user)):
    if not user.is_admin:
        raise HTTPException(status_code=403, detail="Siz admin emassiz")
    return user

@app.get("/api/admin/stats")
async def get_admin_stats(db: Session = Depends(database.get_db), admin: database.User = Depends(get_admin_user)):
    total_users = db.query(database.User).count()
    premium_users = db.query(database.User).filter(database.User.is_premium == True).count()
    
    # Total earnings from payments
    earnings = db.query(database.Payment).filter(database.Payment.status == "success").all()
    total_earnings = sum([p.amount for p in earnings])
    
    # Today's files processed
    today = datetime.utcnow().strftime("%Y-%m-%d")
    today_files = db.query(database.Usage).filter(database.Usage.date == today).all()
    total_files = sum([u.count for u in today_files])
    
    return {
        "total_users": total_users,
        "premium_users": premium_users,
        "total_earnings": total_earnings,
        "today_files": total_files
    }

@app.get("/api/admin/users")
async def get_admin_users(db: Session = Depends(database.get_db), admin: database.User = Depends(get_admin_user)):
    users = db.query(database.User).all()
    return users

@app.post("/api/admin/users/{user_id}/premium")
async def toggle_user_premium(user_id: int, data: dict, db: Session = Depends(database.get_db), admin: database.User = Depends(get_admin_user)):
    user = db.query(database.User).filter(database.User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Foydalanuvchi topilmadi")
    user.is_premium = data.get("is_premium", False)
    db.commit()
    return {"message": "Muvaffaqiyatli o'zgartirildi"}

@app.get("/api/auth/google/callback")
async def google_callback(code: str, db: Session = Depends(database.get_db)):
    """Google shu manzilga qaytaradi. Bunda ham `@app.get` yo'q edi.

    Google Cloud Console'dagi "Authorized redirect URIs" ro'yxatida aynan
    shu manzil turishi kerak:
        https://uzbekpdf.uz/api/auth/google/callback
    """
    token_url = "https://oauth2.googleapis.com/token"
    data = {
        "code": code,
        "client_id": GOOGLE_CLIENT_ID,
        "client_secret": GOOGLE_CLIENT_SECRET,
        "redirect_uri": GOOGLE_REDIRECT_URI,
        "grant_type": "authorization_code",
    }
    async with httpx.AsyncClient() as client:
        response = await client.post(token_url, data=data)
        access_token = response.json().get("access_token")
        
        if not access_token:
            raise HTTPException(status_code=400, detail="Google bilan ulanishda xatolik")
            
        user_info_response = await client.get(
            "https://www.googleapis.com/oauth2/v1/userinfo",
            headers={"Authorization": f"Bearer {access_token}"}
        )
        user_info = user_info_response.json()
    
    email = user_info.get("email")
    google_id = user_info.get("id")
    name = user_info.get("name")
    picture = user_info.get("picture")
    
    db_user = db.query(database.User).filter(database.User.email == email).first()
    if not db_user:
        # Check if username exists, if so append random
        base_username = name.replace(" ", "").lower() if name else email.split("@")[0]
        username = base_username
        while db.query(database.User).filter(database.User.username == username).first():
            import random
            username = f"{base_username}{random.randint(10,999)}"
            
        db_user = database.User(
            username=username,
            email=email,
            google_id=google_id,
            avatar_url=picture
        )
        db.add(db_user)
        db.commit()
        db.refresh(db_user)
        
    access_token_expires = timedelta(minutes=auth.ACCESS_TOKEN_EXPIRE_MINUTES)
    jwt_token = auth.create_access_token(
        data={"sub": db_user.username}, expires_delta=access_token_expires
    )
    
    return RedirectResponse(url=f"/?token={jwt_token}&username={db_user.username}&avatar={picture}")


# --- PDF → Word: "faqat matn" rejimi --------------------------------------
#
# `pdf2docx` sahifani jadval va bo'limlarga bo'lib qayta quradi. Bir ustunli
# oddiy hujjatda bu yaxshi natija beradi, lekin ikki ustunli va rangli fon
# to'rtburchagi bor dizaynda (masalan rezyume) u yo'q jadvalni "topib" oladi:
# matn bir-birining ustiga tushadi va mazmun qirqiladi.
#
# Bu yerdagi yo'l dizaynni umuman takrorlamaydi. Maqsad boshqa: matn to'liq,
# tartibli va Word'da bemalol tahrirlanadigan bo'lsin. Word'ga o'giradigan
# odam odatda aynan shuni xohlaydi.

BULLET_BELGILARI = ("•", "▪", "‣", "·", "–", "—", "-", "*")

# "Ko'rinishni saqlash" rejimi uchun `pdf2docx` sozlamalari.
#
# Standart holatda kutubxona hech qanday chizig'i bo'lmagan joyda ham bo'sh
# joyga qarab jadval "topib" oladi (`parse_stream_table`). Ikki ustunli
# rezyumeda u katak o'lchamlarini noto'g'ri hisoblaydi: matn ustma-ust
# tushadi va mazmun qirqiladi. Aynan shu buzilish kuzatilgan.
#
# Qiymatlar haqiqiy hujjatda o'lchab tanlangan — o'zgartirishdan oldin
# qaytadan o'lchang.
PDF2DOCX_SOZLAMALARI = {
    # Bo'sh joyga qarab jadval yasashni to'xtatadi
    "parse_stream_table": False,
    # Chiziqlari bor haqiqiy jadvallar saqlanadi
    "parse_lattice_table": True,
    # Hoshiyani qirqmaymiz: standart 0.5 matnni sahifa chetiga surib,
    # uzun qatorlarni sig'maydigan qilib qo'yardi
    "page_margin_factor_top": 0.0,
    "page_margin_factor_bottom": 0.0,
    # Bitta sahifadagi xato butun konvertatsiyani to'xtatmasin
    "ignore_page_error": True,
}


def _sahifa_ustunlari(bloklar, sahifa_kengligi):
    """Sahifa ikki ustunlimi? Ha bo'lsa (chap, o'ng) qaytaradi.

    Ustun chegarasi matn umuman tushmagan vertikal yo'lak bo'yicha topiladi.
    Ikki ustunli rezyumeda yon panel va asosiy qism aralashib ketmasligi
    uchun shu kerak — aks holda matn qatorma-qator sakrab o'qib bo'lmas
    holga kelardi.
    """
    if len(bloklar) < 4:
        return None

    KATAK = 100
    band = [False] * KATAK
    for blok in bloklar:
        x0, _, x1, _ = blok["bbox"]
        boshi = max(0, int(x0 / sahifa_kengligi * KATAK))
        oxiri = min(KATAK - 1, int(x1 / sahifa_kengligi * KATAK))
        for i in range(boshi, oxiri + 1):
            band[i] = True

    # Eng keng bo'sh yo'lak. Faqat sahifaning o'rta qismida qidiriladi:
    # chetdagi bo'sh joy hoshiya, u ustun chegarasi emas.
    eng_yaxshi = None
    i = int(KATAK * 0.2)
    while i < int(KATAK * 0.8):
        if band[i]:
            i += 1
            continue
        j = i
        while j < int(KATAK * 0.8) and not band[j]:
            j += 1
        if eng_yaxshi is None or (j - i) > (eng_yaxshi[1] - eng_yaxshi[0]):
            eng_yaxshi = (i, j)
        i = j

    # Tor tirqish oddiy so'z oralig'i bo'lishi mumkin — ustun deb qabul
    # qilmaymiz, aks holda bir ustunli hujjat ham ikkiga bo'linib ketardi.
    if not eng_yaxshi or (eng_yaxshi[1] - eng_yaxshi[0]) < KATAK * 0.04:
        return None

    chegara = (eng_yaxshi[0] + eng_yaxshi[1]) / 2 / KATAK * sahifa_kengligi
    chap = [b for b in bloklar if (b["bbox"][0] + b["bbox"][2]) / 2 < chegara]
    ong = [b for b in bloklar if (b["bbox"][0] + b["bbox"][2]) / 2 >= chegara]

    # Bir tomonda deyarli hech narsa bo'lmasa bu ustun emas, tasodif
    if len(chap) < 2 or len(ong) < 2:
        return None
    return chap, ong


def _asosiy_shrift(hujjat):
    """Hujjatdagi eng ko'p ishlatilgan shrift o'lchami — "oddiy matn" o'lchami.

    Sarlavhani shundan kattaligi bilan ajratamiz. Qat'iy son (masalan 14pt)
    yozib qo'yib bo'lmaydi: har hujjat o'z o'lchamida yoziladi.
    """
    hisob = {}
    for sahifa in hujjat:
        for blok in sahifa.get_text("dict")["blocks"]:
            if blok.get("type") != 0:
                continue
            for qator in blok["lines"]:
                for bolak in qator["spans"]:
                    olcham = round(bolak["size"] * 2) / 2
                    hisob[olcham] = hisob.get(olcham, 0) + len(bolak["text"].strip())
    if not hisob:
        return 11.0
    return max(hisob.items(), key=lambda juft: juft[1])[0]


def _pdf_dan_toza_docx(pdf_yoli, docx_yoli):
    """PDF matnini toza Word hujjatiga yozadi. Xatboshilar sonini qaytaradi.

    0 qaytsa — hujjatda matn yo'q (skanerlangan rasm). Chaqiruvchi buni
    tushunarli xabarga aylantiradi, chunki bo'sh fayl bergandan ko'ra
    sababini aytish yaxshiroq.
    """
    import pymupdf
    from docx import Document
    from docx.shared import Pt

    hujjat = pymupdf.open(pdf_yoli)
    asosiy = _asosiy_shrift(hujjat)
    word = Document()
    yozilgan = 0

    for sahifa in hujjat:
        bloklar = [b for b in sahifa.get_text("dict")["blocks"] if b.get("type") == 0]
        if not bloklar:
            continue

        ustunlar = _sahifa_ustunlari(bloklar, sahifa.rect.width)
        if ustunlar:
            # Avval chap ustun to'liq, keyin o'ng — rezyumedagi tabiiy tartib
            tartib = sorted(ustunlar[0], key=lambda b: b["bbox"][1]) + \
                     sorted(ustunlar[1], key=lambda b: b["bbox"][1])
        else:
            tartib = sorted(bloklar, key=lambda b: b["bbox"][1])

        for blok in tartib:
            blok_ongi = blok["bbox"][2]
            yigilgan, yigilgan_qalin, yigilgan_olcham = "", True, 0.0

            # Blok oqar matnmi yoki ro'yxatmi?
            #
            # Oqar matnda qator sahifa chetiga yetib o'raladi, ya'ni keyingi
            # qator bilan bitta xatboshi. Ro'yxatda esa har qator alohida
            # fikr: "O'zbek — Ona tili" va "Ingliz — B1-B2" qo'shilib
            # ketmasligi kerak.
            #
            # Faqat qatorning o'ng cheti bo'yicha ajratib bo'lmaydi: blokdagi
            # eng uzun qator har doim chetga yetadi va "o'ralgan" bo'lib
            # ko'rinadi. Shuning uchun qator UZUNLIGIGA qaraymiz — ro'yxat
            # bandlari qisqa bo'ladi.
            uzunliklar = [
                len("".join(b["text"] for b in q["spans"]).strip())
                for q in blok["lines"]
            ]
            uzunliklar = [u for u in uzunliklar if u]
            oqar_matn = len(uzunliklar) >= 2 and (
                sorted(uzunliklar)[len(uzunliklar) // 2] >= 45
            )

            def chiqar(matn, qalin, olcham):
                nonlocal yozilgan
                matn = matn.strip()
                if not matn:
                    return
                uslub = None
                belgili = matn[:1] in BULLET_BELGILARI and len(matn) > 2
                if belgili:
                    matn = matn[1:].strip()
                    uslub = "List Bullet"
                elif olcham >= asosiy * 1.45:
                    uslub = "Heading 1"
                elif olcham >= asosiy * 1.15 or (qalin and olcham >= asosiy):
                    uslub = "Heading 2"

                xatboshi = word.add_paragraph(style=uslub) if uslub else word.add_paragraph()
                yugurish = xatboshi.add_run(matn)
                if uslub is None:
                    yugurish.bold = qalin
                    yugurish.font.size = Pt(olcham)
                yozilgan += 1

            for qator in blok["lines"]:
                matn = "".join(bolak["text"] for bolak in qator["spans"])
                if not matn.strip():
                    continue
                olcham = max((b["size"] for b in qator["spans"]), default=asosiy)
                # PyMuPDF bayroqlari: 2**4 — qalin
                qalin = all(b["flags"] & 16 for b in qator["spans"])

                if yigilgan:
                    yigilgan += " " + matn.strip()
                    yigilgan_qalin = yigilgan_qalin and qalin
                    yigilgan_olcham = max(yigilgan_olcham, olcham)
                else:
                    yigilgan, yigilgan_qalin, yigilgan_olcham = matn.strip(), qalin, olcham

                # Oqar matnda qator chetga yetgan bo'lsa — o'ralgan, keyingisi
                # bilan qo'shiladi. Ro'yxatda esa har qator o'zicha qoladi.
                oralgan = oqar_matn and qator["bbox"][2] >= blok_ongi - max(
                    2.0, (blok_ongi - blok["bbox"][0]) * 0.03
                )
                if not oralgan:
                    chiqar(yigilgan, yigilgan_qalin, yigilgan_olcham)
                    yigilgan, yigilgan_qalin, yigilgan_olcham = "", True, 0.0

            chiqar(yigilgan, yigilgan_qalin, yigilgan_olcham)

    hujjat.close()
    if yozilgan:
        word.save(docx_yoli)
    return yozilgan


@app.post("/api/convert/pdf-to-docx")
async def convert_pdf_to_docx(
    request: Request,
    background_tasks: BackgroundTasks, 
    file: UploadFile = File(...),
    mode: str = Form("text"),
    db: Session = Depends(database.get_db),
    user = Depends(get_optional_user)
):
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Faqat PDF fayllar qabul qilinadi")

    # Check limits
    file.file.seek(0, 2)
    fsize = file.file.tell()
    file.file.seek(0)
    await check_usage_limit(fsize, request, db, user)

    unique_id = str(uuid.uuid4())
    pdf_path = os.path.join(TEMP_DIR, f"{unique_id}.pdf")
    docx_path = os.path.join(TEMP_DIR, f"{unique_id}.docx")

    try:
        # Save the uploaded PDF
        with open(pdf_path, "wb") as buffer:
            buffer.write(await file.read())

        if mode == "layout":
            # Ko'rinishni saqlashga urinish. Sozlamalar ataylab berilgan:
            # standart holatda `pdf2docx` bo'sh joyga qarab jadval "ixtiro
            # qiladi" va ikki ustunli dizaynda katak o'lchamini noto'g'ri
            # hisoblab, matnni qirqib qo'yadi.
            cv = Converter(pdf_path)
            cv.convert(docx_path, **PDF2DOCX_SOZLAMALARI)
            cv.close()
        else:
            # Standart yo'l: toza, tahrirlanadigan matn. Hech qachon
            # buzilmaydi, chunki dizaynni takrorlashga urinmaydi.
            if _pdf_dan_toza_docx(pdf_path, docx_path) == 0:
                raise HTTPException(
                    status_code=400,
                    detail="Bu PDF ichida matn yo'q — u skanerlangan rasmdan iborat. "
                           "Uni Word'ga aylantirib bo'lmaydi.",
                )

        # Check if the file was created successfully
        if not os.path.exists(docx_path):
            raise Exception("Word fayli yaratilmadi")

        # Schedule background cleanup after the response is sent
        background_tasks.add_task(cleanup_files, pdf_path, docx_path)

        await increment_usage(request, db, user)

        # Return the file with a safe ASCII filename in the header

        # (The frontend will rename it to the actual original name)
        return FileResponse(
            path=docx_path, 
            filename="converted.docx",
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except HTTPException:
        # Sababi ma'lum bo'lgan xato (masalan skanerlangan PDF) o'z holicha
        # o'tsin — uni 500 ga aylantirsak foydalanuvchi sababni bilmay qolardi.
        cleanup_files(pdf_path, docx_path)
        raise

    except Exception as e:
        # If error occurs, clean up immediately
        cleanup_files(pdf_path, docx_path)
        raise HTTPException(status_code=500, detail=f"Konvertatsiyada xatolik: {str(e)}")

@app.post("/api/convert/merge-pdf")
async def merge_pdf(
    request: Request,
    background_tasks: BackgroundTasks, 
    files: List[UploadFile] = File(...),
    db: Session = Depends(database.get_db),
    user = Depends(get_optional_user)
):
    if not files or len(files) < 2:
        raise HTTPException(status_code=400, detail="Kamida 2 ta PDF fayl kerak")
    
    # Calculate total size and check limits
    total_size = 0
    for f in files:
        f.file.seek(0, 2)
        total_size += f.file.tell()
        f.file.seek(0)
    
    await check_usage_limit(total_size, request, db, user)

    merger = PdfWriter()
    temp_files = []
    
    unique_id = str(uuid.uuid4())
    output_path = os.path.join(TEMP_DIR, f"{unique_id}_merged.pdf")
    
    try:
        for file in files:
            if not file.filename.lower().endswith('.pdf'):
                raise HTTPException(status_code=400, detail="Faqat PDF fayllar qabul qilinadi")
            
            temp_path = os.path.join(TEMP_DIR, f"{uuid.uuid4()}.pdf")
            temp_files.append(temp_path)
            with open(temp_path, "wb") as buffer:
                # OPTIMIZATION: Write in chunks instead of full read
                while True:
                    chunk = await file.read(1024 * 1024) # 1MB chunks
                    if not chunk: break
                    buffer.write(chunk)
            merger.append(temp_path)

            
        merger.write(output_path)
        merger.close()
        
        await increment_usage(request, db, user)
        
        background_tasks.add_task(cleanup_files, output_path, *temp_files)

        return FileResponse(path=output_path, filename="uzbekpdf_merged.pdf", media_type='application/pdf')
    except Exception as e:
        cleanup_files(output_path, *temp_files)
        raise HTTPException(status_code=500, detail=f"Birlashtirishda xatolik yuz berdi: {str(e)}")

@app.post("/api/convert/split-pdf")
async def split_pdf(
    request: Request,
    background_tasks: BackgroundTasks, 
    file: UploadFile = File(...),
    db: Session = Depends(database.get_db),
    user = Depends(get_optional_user)
):
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Faqat PDF fayl qabul qilinadi")
    
    file.file.seek(0, 2)
    fsize = file.file.tell()
    file.file.seek(0)
    await check_usage_limit(fsize, request, db, user)
        
    unique_id = str(uuid.uuid4())
    pdf_path = os.path.join(TEMP_DIR, f"{unique_id}.pdf")
    zip_path = os.path.join(TEMP_DIR, f"{unique_id}_split.zip")
    
    try:
        await save_upload_file(file, pdf_path)
            
        reader = PdfReader(pdf_path)

        
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for i, page in enumerate(reader.pages):
                writer = PdfWriter()
                writer.add_page(page)
                page_path = os.path.join(TEMP_DIR, f"page_{i+1}.pdf")
                with open(page_path, "wb") as f:
                    writer.write(f)
                zipf.write(page_path, f"page_{i+1}.pdf")
                os.remove(page_path)
                
        background_tasks.add_task(cleanup_files, pdf_path, zip_path)
        return FileResponse(path=zip_path, filename="uzbekpdf_split.zip", media_type='application/zip')
    except Exception as e:
        cleanup_files(pdf_path, zip_path)
        raise HTTPException(status_code=500, detail=f"Kesishda xatolik yuz berdi: {str(e)}")

@app.post("/api/convert/pdf-to-jpg")
async def pdf_to_jpg(
    background_tasks: BackgroundTasks, 
    file: UploadFile = File(...),
    current_user: database.User = Depends(auth.get_current_user)
):
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Faqat PDF fayl qabul qilinadi")
        
    unique_id = str(uuid.uuid4())
    pdf_path = os.path.join(TEMP_DIR, f"{unique_id}.pdf")
    zip_path = os.path.join(TEMP_DIR, f"{unique_id}_images.zip")
    
    try:
        with open(pdf_path, "wb") as buffer:
            buffer.write(await file.read())
            
        images = convert_from_path(pdf_path, dpi=200)
        
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for i, image in enumerate(images):
                img_path = os.path.join(TEMP_DIR, f"page_{i+1}.jpg")
                image.save(img_path, "JPEG")
                zipf.write(img_path, f"page_{i+1}.jpg")
                os.remove(img_path)
                
        background_tasks.add_task(cleanup_files, pdf_path, zip_path)
        return FileResponse(path=zip_path, filename="uzbekpdf_images.zip", media_type='application/zip')
    except Exception as e:
        cleanup_files(pdf_path, zip_path)
        raise HTTPException(status_code=500, detail=f"Rasmlarga aylantirishda xatolik: {str(e)}")

@app.post("/api/convert/protect-pdf")
async def protect_pdf(
    request: Request,
    background_tasks: BackgroundTasks, 
    file: UploadFile = File(...),
    password: str = Form(...),
    db: Session = Depends(database.get_db),
    user = Depends(get_optional_user)
):
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Faqat PDF fayl qabul qilinadi")
    
    file.file.seek(0, 2)
    fsize = file.file.tell()
    file.file.seek(0)
    await check_usage_limit(fsize, request, db, user)

    unique_id = str(uuid.uuid4())
    pdf_path = os.path.join(TEMP_DIR, f"{unique_id}.pdf")
    output_path = os.path.join(TEMP_DIR, f"{unique_id}_protected.pdf")
    
    try:
        await save_upload_file(file, pdf_path)

            
        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        
        for page in reader.pages:
            writer.add_page(page)
            
        writer.encrypt(password)
        
        with open(output_path, "wb") as f:
            writer.write(f)
            
        background_tasks.add_task(cleanup_files, pdf_path, output_path)
        return FileResponse(path=output_path, filename="uzbekpdf_protected.pdf", media_type='application/pdf')
    except Exception as e:
        cleanup_files(pdf_path, output_path)
        raise HTTPException(status_code=500, detail=f"Parol o'rnatishda xatolik: {str(e)}")

@app.post("/api/convert/unlock-pdf")
async def unlock_pdf(
    request: Request,
    background_tasks: BackgroundTasks, 
    file: UploadFile = File(...),
    password: str = Form(...),
    db: Session = Depends(database.get_db),
    user = Depends(get_optional_user)
):
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Faqat PDF fayl qabul qilinadi")
    
    file.file.seek(0, 2)
    fsize = file.file.tell()
    file.file.seek(0)
    await check_usage_limit(fsize, request, db, user)
        
    unique_id = str(uuid.uuid4())
    pdf_path = os.path.join(TEMP_DIR, f"{unique_id}.pdf")
    output_path = os.path.join(TEMP_DIR, f"{unique_id}_unlocked.pdf")
    
    try:
        await save_upload_file(file, pdf_path)

            
        reader = PdfReader(pdf_path)
        if reader.is_encrypted:
            success = reader.decrypt(password)
            if not success:
                raise Exception("Noto'g'ri parol!")
                
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
            
        with open(output_path, "wb") as f:
            writer.write(f)
            
        await increment_usage(request, db, user)
        background_tasks.add_task(cleanup_files, pdf_path, output_path)
        return FileResponse(path=output_path, filename="uzbekpdf_unlocked.pdf", media_type='application/pdf')

    except Exception as e:
        cleanup_files(pdf_path, output_path)
        raise HTTPException(status_code=400, detail=f"Parolni yechishda xatolik: {str(e)}")

@app.post("/api/convert/rotate-pdf")
async def rotate_pdf(
    request: Request,
    background_tasks: BackgroundTasks, 
    file: UploadFile = File(...),
    angle: int = Form(...),
    db: Session = Depends(database.get_db),
    user = Depends(get_optional_user)
):
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Faqat PDF fayl qabul qilinadi")
    
    file.file.seek(0, 2)
    fsize = file.file.tell()
    file.file.seek(0)
    await check_usage_limit(fsize, request, db, user)
        
    unique_id = str(uuid.uuid4())
    pdf_path = os.path.join(TEMP_DIR, f"{unique_id}.pdf")
    output_path = os.path.join(TEMP_DIR, f"{unique_id}_rotated.pdf")
    
    try:
        await save_upload_file(file, pdf_path)

            
        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        
        for page in reader.pages:
            writer.add_page(page.rotate(angle))
            
        with open(output_path, "wb") as f:
            writer.write(f)
            
        await increment_usage(request, db, user)
        background_tasks.add_task(cleanup_files, pdf_path, output_path)
        return FileResponse(path=output_path, filename="uzbekpdf_rotated.pdf", media_type='application/pdf')

    except Exception as e:
        cleanup_files(pdf_path, output_path)
        raise HTTPException(status_code=500, detail=f"Burishda xatolik: {str(e)}")

@app.post("/api/convert/compress-pdf")
async def compress_pdf(
    request: Request,
    background_tasks: BackgroundTasks, 
    file: UploadFile = File(...),
    db: Session = Depends(database.get_db),
    user = Depends(get_optional_user)
):
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Faqat PDF fayl qabul qilinadi")
    
    file.file.seek(0, 2)
    fsize = file.file.tell()
    file.file.seek(0)
    await check_usage_limit(fsize, request, db, user)
        
    unique_id = str(uuid.uuid4())
    pdf_path = os.path.join(TEMP_DIR, f"{unique_id}.pdf")
    output_path = os.path.join(TEMP_DIR, f"{unique_id}_compressed.pdf")
    
    try:
        await save_upload_file(file, pdf_path)
            
        # Ghostscript command for compression
        # /screen (lowest quality, smallest size), /ebook (medium), /printer (high), /prepress (highest)
        # We use /ebook as a good balance for compression
        cmd = [
            "/usr/bin/gs", "-sDEVICE=pdfwrite", "-dCompatibilityLevel=1.4",
            "-dPDFSETTINGS=/ebook", "-dNOPAUSE", "-dQUIET", "-dBATCH",
            f"-sOutputFile={output_path}", pdf_path
        ]

        # Ensure full PATH is available for subprocesses
        env = os.environ.copy()
        env["PATH"] = "/usr/local/sbin:/usr/local/bin:/usr/sbin:/usr/bin:/sbin:/bin:" + env.get("PATH", "")
        
        process = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, env=env)

        if process.returncode != 0:
            raise Exception("Ghostscript xatosi yoki o'rnatilmagan")
            
        if not os.path.exists(output_path):
            raise Exception("Kichraytirilgan fayl yaratilmadi")
            
        await increment_usage(request, db, user)
        background_tasks.add_task(cleanup_files, pdf_path, output_path)
        return FileResponse(path=output_path, filename="uzbekpdf_compressed.pdf", media_type='application/pdf')

    except Exception as e:
        cleanup_files(pdf_path, output_path)
        raise HTTPException(status_code=500, detail=f"Kichraytirishda xatolik yuz berdi: {str(e)}")


@app.post("/api/convert/office-to-pdf")
async def office_to_pdf(
    request: Request,
    background_tasks: BackgroundTasks, 
    file: UploadFile = File(...),
    db: Session = Depends(database.get_db),
    user = Depends(get_optional_user)
):
    valid_extensions = ('.docx', '.doc', '.pptx', '.ppt', '.xlsx', '.xls')
    if not file.filename.lower().endswith(valid_extensions):
        raise HTTPException(status_code=400, detail="Faqat Word, Excel yoki PowerPoint fayllari qabul qilinadi")
    
    file.file.seek(0, 2)
    fsize = file.file.tell()
    file.file.seek(0)
    await check_usage_limit(fsize, request, db, user)
        
    unique_id = str(uuid.uuid4())
    ext = os.path.splitext(file.filename)[1].lower()
    input_path = os.path.join(TEMP_DIR, f"{unique_id}{ext}")
    temp_files = [input_path]
    
    # LibreOffice automatically names the output based on input filename
    output_path = os.path.join(TEMP_DIR, f"{unique_id}.pdf")
    
    try:
        await save_upload_file(file, input_path)
            
        target_path = input_path
        
        # If it's Excel, we optimize it for printing
        if ext in ['.xlsx', '.xls']:
            if ext == '.xls':
                # Convert .xls to .xlsx first using LibreOffice
                cmd_xls = [
                    "/usr/bin/libreoffice", "-env:UserInstallation=file:///tmp/libo_pdf_converter",
                    "--headless", "--convert-to", "xlsx", 
                    "--outdir", TEMP_DIR, input_path
                ]
                env = os.environ.copy()
                env["PATH"] = "/usr/local/sbin:/usr/local/bin:/usr/sbin:/usr/bin:/sbin:/bin:" + env.get("PATH", "")
                subprocess.run(cmd_xls, stdout=subprocess.PIPE, stderr=subprocess.PIPE, env=env)

                xlsx_path = os.path.join(TEMP_DIR, f"{unique_id}.xlsx")
                if os.path.exists(xlsx_path):
                    target_path = xlsx_path
                    temp_files.append(xlsx_path)
            
            # Apply OpenPyXL scaling
            try:
                wb = openpyxl.load_workbook(target_path)
                for ws in wb.worksheets:
                    ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
                    ws.page_setup.paperSize = ws.PAPERSIZE_A4
                    ws.page_setup.fitToPage = True
                    ws.page_setup.fitToHeight = False
                    ws.page_setup.fitToWidth = 1
                    
                    ws.page_margins.left = 0.25
                    ws.page_margins.right = 0.25
                    ws.page_margins.top = 0.75
                    ws.page_margins.bottom = 0.75
                
                # Save as a new optimized file
                optimized_path = os.path.join(TEMP_DIR, f"{unique_id}_optimized.xlsx")
                wb.save(optimized_path)
                target_path = optimized_path
                temp_files.append(optimized_path)
            except Exception as e:
                # If optimization fails, just continue with the unoptimized file
                print(f"Excel optimization failed: {e}")
                
            # Update expected output path since the base name changed
            base_name = os.path.splitext(os.path.basename(target_path))[0]
            output_path = os.path.join(TEMP_DIR, f"{base_name}.pdf")
            temp_files.append(output_path)
            
        cmd = [
            "/usr/bin/libreoffice", "-env:UserInstallation=file:///tmp/libo_pdf_converter",
            "--headless", "--convert-to", "pdf", 
            "--outdir", TEMP_DIR, target_path
        ]
        
        env = os.environ.copy()
        env["PATH"] = "/usr/local/sbin:/usr/local/bin:/usr/sbin:/usr/bin:/sbin:/bin:" + env.get("PATH", "")
        process = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, env=env)

        if process.returncode != 0:
            error_msg = process.stderr.decode()
            raise Exception(f"LibreOffice xatosi: {error_msg}")
            
        if not os.path.exists(output_path):
            raise Exception("PDF fayl yaratilmadi")
            
        await increment_usage(request, db, user)
        background_tasks.add_task(cleanup_files, *temp_files)
        return FileResponse(path=output_path, filename="uzbekpdf_office_result.pdf", media_type='application/pdf')
    except Exception as e:
        cleanup_files(*temp_files)
        raise HTTPException(status_code=500, detail=f"PDF ga o'tkazishda xatolik: {str(e)}")


@app.post("/api/convert/watermark")
async def watermark_pdf(
    background_tasks: BackgroundTasks, 
    file: UploadFile = File(...),
    watermark_text: str = Form(...),
    current_user: database.User = Depends(auth.get_current_user)
):
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Faqat PDF fayl qabul qilinadi")
        
    unique_id = str(uuid.uuid4())
    pdf_path = os.path.join(TEMP_DIR, f"{unique_id}.pdf")
    watermark_path = os.path.join(TEMP_DIR, f"{unique_id}_wm.pdf")
    output_path = os.path.join(TEMP_DIR, f"{unique_id}_watermarked.pdf")
    
    try:
        with open(pdf_path, "wb") as buffer:
            buffer.write(await file.read())
            
        # Create Watermark PDF
        c = canvas.Canvas(watermark_path, pagesize=letter)
        c.setFont("Helvetica-Bold", 60)
        c.setFillColor(Color(0.5, 0.5, 0.5, alpha=0.3)) # Light gray with opacity
        c.translate(200, 200)
        c.rotate(45)
        c.drawCentredString(200, 0, watermark_text)
        c.save()
        
        watermark_reader = PdfReader(watermark_path)
        watermark_page = watermark_reader.pages[0]
        
        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        
        for page in reader.pages:
            page.merge_page(watermark_page)
            writer.add_page(page)
            
        with open(output_path, "wb") as f:
            writer.write(f)
            
        background_tasks.add_task(cleanup_files, pdf_path, watermark_path, output_path)
        return FileResponse(path=output_path, filename="uzbekpdf_watermarked.pdf", media_type='application/pdf')
    except Exception as e:
        cleanup_files(pdf_path, watermark_path, output_path)
        raise HTTPException(status_code=500, detail=f"Belgi tushirishda xatolik: {str(e)}")

@app.post("/api/convert/edit-pdf")
async def edit_pdf(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    annotations_json: str = Form(...),
    page_order_json: str = Form(...),
    current_user: database.User = Depends(auth.get_current_user)
):
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Faqat PDF fayl qabul qilinadi")
    
    unique_id = str(uuid.uuid4())
    input_path = os.path.join(TEMP_DIR, f"{unique_id}_input.pdf")
    output_path = os.path.join(TEMP_DIR, f"{unique_id}_edited.pdf")
    
    try:
        with open(input_path, "wb") as buffer:
            buffer.write(await file.read())
        
        annotations = json.loads(annotations_json)  # {"1": "data:image/png;base64,...", ...}
        page_order = json.loads(page_order_json)    # [3, 1, 4] vs
        
        reader = PdfReader(input_path)
        writer = PdfWriter()
        
        for orig_num in page_order:
            if isinstance(orig_num, str) and orig_num.startswith('blank_'):
                parts = orig_num.split(':')
                page_num_str = parts[0]
                if len(parts) >= 3:
                    w = float(parts[1])
                    h = float(parts[2])
                else:
                    w = float(reader.pages[0].mediabox.width) if reader.pages else 595.27
                    h = float(reader.pages[0].mediabox.height) if reader.pages else 841.89
                page = PageObject.create_blank_page(width=w, height=h)
            else:
                page_idx = int(orig_num) - 1
                if page_idx < 0 or page_idx >= len(reader.pages):
                    continue
                page = reader.pages[page_idx]
                page_num_str = str(orig_num)
            
            if page_num_str in annotations:
                img_data_url = annotations[page_num_str]
                img_b64 = img_data_url.split(',', 1)[1]
                img_bytes = base64.b64decode(img_b64)
                
                # Get page size in PDF points
                page_width  = float(page.mediabox.width)
                page_height = float(page.mediabox.height)
                
                # Build a single-page overlay PDF from the annotation PNG
                overlay_buf = io.BytesIO()
                c = canvas.Canvas(overlay_buf, pagesize=(page_width, page_height))
                pil_img = PILImage.open(io.BytesIO(img_bytes)).convert('RGBA')
                img_reader = ImageReader(pil_img)
                c.drawImage(img_reader, 0, 0,
                            width=page_width, height=page_height,
                            mask='auto', preserveAspectRatio=False)
                c.save()
                overlay_buf.seek(0)
                
                # Merge overlay onto original page
                overlay_reader = PdfReader(overlay_buf)
                page.merge_page(overlay_reader.pages[0])
            
            writer.add_page(page)
        
        with open(output_path, "wb") as f:
            writer.write(f)
        
        background_tasks.add_task(cleanup_files, input_path, output_path)
        return FileResponse(path=output_path, filename="uzbekpdf_edited.pdf", media_type="application/pdf")
    
    except Exception as e:
        cleanup_files(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Tahrirda xatolik: {str(e)}")

# This endpoint serves index.html if we want to run both together
@app.get("/")
async def root():
    try:
        with open("index.html", "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        return HTMLResponse(content="<h1>index.html topilmadi</h1>")

@app.get("/login")
async def login_page():
    """Kirish sahifasi.

    `login.html` fayli bor edi, lekin unga marshrut yozilmagan — ya'ni
    sahifa umuman ochilmasdi. Barcha sahifalar esa foydalanuvchini aynan
    `/login` ga yuboradi (`index.html`, `editor.html`, `profile.html`,
    `premium.html`, `admin.html`), demak kirish butun saytda ishlamasdi.

    Ro'yxatdan o'tish ham shu sahifada: forma `/api/register` ga yuboradi,
    shuning uchun alohida `/register` marshruti kerak emas.
    """
    try:
        with open("login.html", "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        return HTMLResponse(content="<h1>login.html topilmadi</h1>")


@app.get("/profile")
async def profile_page():
    try:
        with open("profile.html", "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        return HTMLResponse(content="<h1>profile.html topilmadi</h1>")

# Admin paneli hozircha yopiq.
#
# Ikki sabab bilan olib tashlandi:
#   * bazada birorta ham admin yo'q (`is_admin`), ya'ni panel hech kimga
#     ochilmaydi va faqat bo'sh raqamlar ko'rsatadi;
#   * sahifa parolsiz berilardi — ma'lumot chiqmasa ham (API `get_admin_user`
#     bilan himoyalangan), begona odam admin panelining o'zini ko'ra olardi.
#
# `admin.html` va `/api/admin/*` marshrutlari joyida qoldi. Qaytarish uchun:
#   1. o'z hisobingizga admin huquqini bering:
#         UPDATE users SET is_admin = 1 WHERE email = '...';
#   2. quyidagi marshrutni tiklang;
#   3. `admin.html` ga tekshiruv qo'shing: kirmagan yoki admin bo'lmagan
#      odamni `/login` ga qaytarsin — hozir u shunchaki bo'sh sahifa
#      ko'rsatadi.
#
# @app.get("/admin")
# async def admin_page():
#     with open("admin.html", "r", encoding="utf-8") as f:
#         return HTMLResponse(content=f.read())


from datetime import datetime


@app.post("/api/pay/create")
async def create_payment(plan: str = Form(...), db: Session = Depends(database.get_db), current_user: database.User = Depends(auth.get_current_user)):
    amounts = {
        "monthly": 20000,
        "yearly": 100000
    }
    amount = amounts.get(plan, 0)
    if amount == 0:
        raise HTTPException(status_code=400, detail="Noto'g'ri reja tanlandi")

    # Create local payment record
    new_pay = database.Payment(
        user_id=current_user.id,
        amount=amount,
        plan=plan,
        status="pending",
        created_at=datetime.utcnow().isoformat()
    )
    db.add(new_pay)
    db.commit()
    db.refresh(new_pay)

    # Payme Config (Dummy values for now)
    MERCHANT_ID = "65e6d6c29b9f7a7a2a7a2a7a" # O'zingizning Merchant ID ni qo'ying
    
    # Generate Payme Link
    # Payme expects params in base64: m=...;ac.user_id=...;a=...
    params = f"m={MERCHANT_ID};ac.user_id={current_user.id};a={amount * 100}"
    import base64
    b64_params = base64.b64encode(params.encode()).decode()
    pay_url = f"https://checkout.paycom.uz/{b64_params}"

    return {"pay_url": pay_url}

@app.get("/premium")

async def premium_page():
    try:
        with open("premium.html", "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        return HTMLResponse(content="<h1>premium.html topilmadi</h1>")

@app.get("/editor")
async def editor_page():
    try:
        with open("editor.html", "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        return HTMLResponse(content="<h1>editor.html topilmadi</h1>")

@app.get("/about")
async def about_page():
    try:
        with open("about.html", "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        return HTMLResponse(content="<h1>about.html topilmadi</h1>")

@app.get("/contact")
async def contact_page():
    try:
        with open("contact.html", "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        return HTMLResponse(content="<h1>contact.html topilmadi</h1>")



